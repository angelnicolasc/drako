"""DOCX generator for EU AI Act articles.

The mandatory legal disclaimer is the first paragraph of every document
and is not removable: it is sourced from `drako.comply.LEGAL_DISCLAIMER`
and prepended unconditionally before any article-specific content.
"""

from __future__ import annotations

import re
from pathlib import Path

from drako.comply import LEGAL_DISCLAIMER, ComplianceContext


def _safe_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]", "_", name) or "system"


def _import_docx():
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ImportError as exc:
        raise ImportError(
            "drako[comply] extras are required for DOCX generation. "
            "Install with: pip install drako[comply]"
        ) from exc
    return Document, Pt


def _render_article(doc, article: int, ctx: ComplianceContext) -> None:  # noqa: ANN001
    if article == 9:
        _render_article_9(doc, ctx)
    elif article == 11:
        _render_article_11(doc, ctx)
    elif article == 12:
        _render_article_12(doc, ctx)
    elif article == 14:
        _render_article_14(doc, ctx)
    else:
        raise ValueError(f"Unknown article: {article}")


def _heading(doc, text: str, level: int = 1) -> None:  # noqa: ANN001
    doc.add_heading(text, level=level)


def _para(doc, text: str) -> None:  # noqa: ANN001
    doc.add_paragraph(text or "—")


def _bullets(doc, items: list[str]) -> None:  # noqa: ANN001
    if not items:
        doc.add_paragraph("—")
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _render_article_9(doc, ctx: ComplianceContext) -> None:  # noqa: ANN001
    _heading(doc, "1. Risk management system")
    _para(doc, ctx.system_purpose)
    _heading(doc, "2. Identified risks", level=2)
    _bullets(doc, [f"{f.get('policy_id')}: {f.get('message')}" for f in ctx.risk_findings] or ["No findings recorded."])
    _heading(doc, "3. Residual risks", level=2)
    _bullets(doc, ctx.residual_risks)
    _heading(doc, "4. Mitigation measures", level=2)
    _para(doc, f"Governance score: {ctx.governance_score}/100. Determinism: {ctx.determinism_score}/100.")
    _heading(doc, "5. Post-market monitoring", level=2)
    _para(doc, ctx.update_cadence)


def _render_article_11(doc, ctx: ComplianceContext) -> None:  # noqa: ANN001
    _heading(doc, "1. System description")
    _para(doc, ctx.system_purpose)
    _heading(doc, "2. Architecture", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Component", "Type", "Detail"
    for a in ctx.agents:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = str(a.get("name", "")), "Agent", str(a.get("framework", ""))
    for t in ctx.tools:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = str(t.get("name", "")), "Tool", ""
    for m in ctx.models:
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = str(m.get("name", "")), "Model", ""
    _heading(doc, "3. Training data", level=2)
    _bullets(doc, ctx.training_data_sources)
    _heading(doc, "4. Known limitations", level=2)
    _bullets(doc, ctx.known_limitations)


def _render_article_12(doc, ctx: ComplianceContext) -> None:  # noqa: ANN001
    _heading(doc, "1. Logging system")
    _para(doc, f"Audit trail enabled: {'yes' if ctx.audit_trail_enabled else 'no'}")
    _heading(doc, "2. Retention period", level=2)
    _para(doc, ctx.data_retention_policy)
    _heading(doc, "3. Access controls", level=2)
    _para(doc, ctx.escalation_path)


def _render_article_14(doc, ctx: ComplianceContext) -> None:  # noqa: ANN001
    _heading(doc, "1. Oversight measures")
    _para(doc, f"{ctx.oversight_responsible_name} ({ctx.oversight_responsible_role})")
    _heading(doc, "2. HITL configuration", level=2)
    _para(doc, f"HITL configured: {'yes' if ctx.hitl_configured else 'no'}")
    _bullets(doc, ctx.hitl_triggers)
    _heading(doc, "3. Intervention procedures", level=2)
    _para(doc, ctx.incident_response_procedure)
    _heading(doc, "4. Out-of-scope use cases", level=2)
    _bullets(doc, ctx.out_of_scope_use_cases)


def generate_docx(article: int, ctx: ComplianceContext, output_dir: Path) -> Path:
    """Generate `article_{N}_{system_name}.docx` and return its path."""
    from drako.comply.generators import ARTICLE_META

    Document, _Pt = _import_docx()
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # Disclaimer is mandatory and always the first paragraph.
    disclaimer_para = doc.add_paragraph()
    run = disclaimer_para.add_run(LEGAL_DISCLAIMER)
    run.bold = True

    title = ARTICLE_META[article]["title"]
    doc.add_heading(title, level=0)
    doc.add_paragraph(
        f"System: {ctx.system_name} ({ctx.system_version})  ·  "
        f"Provider: {ctx.provider_entity}  ·  "
        f"Generated: {ctx.generated_at.strftime('%Y-%m-%d %H:%M UTC')}  ·  "
        f"Drako v{ctx.drako_version}"
    )

    _render_article(doc, article, ctx)

    fname = f"article_{article}_{_safe_filename(ctx.system_name)}.docx"
    out_path = output_dir / fname
    doc.save(str(out_path))
    return out_path
